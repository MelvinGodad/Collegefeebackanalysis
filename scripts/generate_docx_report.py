from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_report():
    document = Document()

    # Title
    document.add_heading('ADvantage Dashboard', 0)
    
    subtitle = document.add_paragraph('Marketing Intelligence & Campaign Analytics Report')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    document.add_paragraph('\n\n\n')
    p = document.add_paragraph('Prepared By: Melvin Godad')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_page_break()

    # Executive Summary
    document.add_heading('1. Executive Summary', level=1)
    document.add_paragraph(
        'ADvantage is a next-generation Marketing Intelligence Dashboard designed to unify fragmented advertising data into a single source of truth. '
        'In an era where digital campaigns span primarily Facebook, Instagram, and Google Ads, marketers struggle to consolidate performance metrics manually.'
    )
    document.add_paragraph(
        'This project solves that challenge by providing an automated, visual interface where campaign data—Impressions, Clicks, Conversions, and ROI—is aggregated, '
        'analyzed, and presented in real-time. It enables data-driven decision-making to optimize ad spend and maximize campaign effectiveness.'
    )

    # Objectives
    document.add_heading('2. Project Objectives', level=1)
    objectives = [
        'Unified Visibility: Consolidate data from disparate marketing channels (Social Media, Search, Display) into one view.',
        'Performance Tracking: Monitor real-time KPIs like Click-Through Rate (CTR) and Conversion Rate.',
        'Trend Analysis: Identify high-performing campaigns and scale successful strategies instantly.',
        'Cost Efficiency: Reduce wasted ad spend by highlighting underperforming assets.'
    ]
    for obj in objectives:
        document.add_paragraph(obj, style='List Bullet')

    document.add_page_break()

    # Tech Stack
    document.add_heading('3. Technology Stack & Libraries', level=1)
    document.add_paragraph(
        'The ADvantage platform leverages a modern, lightweight, and high-performance technology stack tailored for speed and visual fidelity.'
    )

    document.add_heading('3.1 Frontend Framework (Web)', level=2)
    document.add_paragraph('HTML5: Semantic markup ensures accessibility and SEO-friendly structure.')
    document.add_paragraph('Tailwind CSS (CDN): utility-first CSS framework used for rapid UI development.')
    document.add_paragraph('Font Awesome (CDN): Utilized for scalable vector icons.')

    document.add_heading('3.2 Visualization & Logic', level=2)
    document.add_paragraph('Chart.js: The core visualization engine (Line Charts, Doughnut Charts, Bar Charts).')
    document.add_paragraph('JavaScript (ES6): Handles data processing on the client side.')

    document.add_heading('3.3 Backend Processing', level=2)
    document.add_paragraph('Python Scripts: Background scripts handle simulation, aggregation, and JSON export.')

    document.add_page_break()

    # System Features
    document.add_heading('4. System Features', level=1)
    
    document.add_heading('4.1 Cross-Platform Integration', level=2)
    document.add_paragraph(
        'ADvantage helps marketers break down silos. By verifying data from Facebook, Instagram, and other digital channels, '
        'it provides a holistic view of the marketing ecosystem.'
    )

    document.add_heading('4.2 Real-Time KPI Cards', level=2)
    document.add_paragraph('Heads-Up displays for critical metrics:')
    document.add_paragraph('Total Reach: Cumulative audience size.', style='List Bullet')
    document.add_paragraph('Engagement Rate: Effectiveness of audience resonance.', style='List Bullet')
    document.add_paragraph('Conversion Cost: Cost Per Acquisition (CPA) tracking.', style='List Bullet')

    document.add_heading('4.3 Smart Landing Page', level=2)
    document.add_paragraph('A dedicated landing page styled with Tailwind CSS to effectively communicate the value proposition.')

    # Conclusion
    document.add_heading('5. Conclusion', level=1)
    document.add_paragraph(
        'ADvantage represents a significant step forward in personal marketing analytics. '
        'By abstracting the complexity of raw data into a clean, visual, and interactive interface, it allows marketers to focus less on spreadsheets and more on strategy.'
    )
    
    document.add_paragraph(
        'The use of Tailwind CSS ensures the application is modern and maintainable, while Chart.js provides the analytical depth required for professional environments.'
    )

    # Save
    document.save('d:/CampusLifeAnalysis/ADvantage_Technical_Report.docx')
    print("Report saved to d:/CampusLifeAnalysis/ADvantage_Technical_Report.docx")

if __name__ == "__main__":
    create_report()
